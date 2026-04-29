import os
import re
from pathlib import Path
import requests
from bs4 import BeautifulSoup
from html2text import HTML2Text
from PIL import Image
import pytesseract
from PyPDF2 import PdfReader
import pdfplumber  # better for complex PDFs
from docx import Document

class TextExtractor:
    @staticmethod
    def from_pdf(file_path):
        """Extract text from PDF (both text-based and simple OCR fallback)"""
        text = ""
        # Try pdfplumber first (better for tables)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except:
            # Fallback to PyPDF2
            try:
                reader = PdfReader(file_path)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            except:
                pass
        return text.strip()

    @staticmethod
    def from_docx(file_path):
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text += "\n" + row_text
        return text.strip()

    @staticmethod
    def from_markdown(file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    @staticmethod
    def from_txt(file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    @staticmethod
    def from_html(file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
            # Remove script/style tags
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text()
            # Clean up whitespace
            text = re.sub(r'\n\s+\n', '\n\n', text)
            return text.strip()

    @staticmethod
    def from_image(file_path):
        """Extract text from images using Tesseract OCR"""
        try:
            img = Image.open(file_path)
            # Basic preprocessing could be added here
            text = pytesseract.image_to_string(img)
            return text.strip()
        except Exception as e:
            if "tesseract is not installed" in str(e).lower() or "no such file" in str(e).lower():
                 return "[OCR Error] Tesseract-OCR is not installed on the system. Please install it to index images."
            return f"[OCR Error] Could not process image: {e}"

    @staticmethod
    def from_url(url):
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        # Extract main content (simplistic – can use readability later)
        for script in soup(["script", "style"]):
            script.decompose()
        text = soup.get_text()
        text = re.sub(r'\n\s+\n', '\n\n', text)
        # Convert to markdown using html2text
        converter = HTML2Text()
        converter.ignore_links = False
        markdown = converter.handle(response.text)
        return markdown

    @staticmethod
    def extract(file_path_or_url):
        """Main entry point – detects type and extracts text"""
        # If it's a URL
        if str(file_path_or_url).startswith(('http://', 'https://')):
            text = TextExtractor.from_url(file_path_or_url)
            return text, {'source_type': 'url', 'mime_type': 'text/html'}

        path = Path(file_path_or_url)
        suffix = path.suffix.lower()

        if suffix == '.pdf':
            text = TextExtractor.from_pdf(path)
            return text, {'source_type': 'pdf', 'mime_type': 'application/pdf'}
        elif suffix == '.docx':
            text = TextExtractor.from_docx(path)
            return text, {'source_type': 'docx', 'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'}
        elif suffix == '.md':
            text = TextExtractor.from_markdown(path)
            return text, {'source_type': 'markdown', 'mime_type': 'text/markdown'}
        elif suffix == '.txt':
            text = TextExtractor.from_txt(path)
            return text, {'source_type': 'text', 'mime_type': 'text/plain'}
        elif suffix in ['.html', '.htm']:
            text = TextExtractor.from_html(path)
            return text, {'source_type': 'html', 'mime_type': 'text/html'}
        elif suffix in ['.png', '.jpg', '.jpeg', '.webp', '.tiff', '.bmp']:
            text = TextExtractor.from_image(path)
            return text, {'source_type': 'image', 'mime_type': f'image/{suffix[1:]}'}
        else:
            raise ValueError(f"Unsupported file type: {suffix}")